"""Tests for the DOCX tool."""

from __future__ import annotations

from unittest.mock import patch

import pytest

from anteroom.tools.office_docx import _MAX_CONTENT_BLOCKS, _MAX_EDIT_OPS, AVAILABLE, DEFINITION, handle

_needs_docx = pytest.mark.skipif(not AVAILABLE, reason="requires python-docx: pip install anteroom[office]")


@pytest.fixture(autouse=True)
def _set_working_dir(tmp_path, monkeypatch):
    monkeypatch.chdir(tmp_path)
    yield


class TestDefinition:
    def test_name(self):
        assert DEFINITION["name"] == "docx"

    def test_required_params(self):
        assert "action" in DEFINITION["parameters"]["required"]
        assert "path" in DEFINITION["parameters"]["required"]


@_needs_docx
class TestCreate:
    @pytest.mark.asyncio
    async def test_create_simple(self, tmp_path):
        result = await handle(
            action="create",
            path="test.docx",
            content_blocks=[
                {"type": "heading", "text": "Title", "level": 1},
                {"type": "paragraph", "text": "Hello world"},
            ],
        )
        assert "error" not in result
        assert result["blocks_written"] == 2
        assert (tmp_path / "test.docx").exists()

    @pytest.mark.asyncio
    async def test_create_with_table(self, tmp_path):
        result = await handle(
            action="create",
            path="tables.docx",
            content_blocks=[
                {"type": "table", "rows": [["A", "B"], ["1", "2"]]},
            ],
        )
        assert "error" not in result
        assert (tmp_path / "tables.docx").exists()

    @pytest.mark.asyncio
    async def test_create_no_blocks(self):
        result = await handle(action="create", path="empty.docx")
        assert "error" in result

    @pytest.mark.asyncio
    async def test_create_too_many_blocks(self):
        blocks = [{"type": "paragraph", "text": f"p{i}"} for i in range(_MAX_CONTENT_BLOCKS + 1)]
        result = await handle(action="create", path="big.docx", content_blocks=blocks)
        assert "error" in result
        assert "Too many" in result["error"]

    @pytest.mark.asyncio
    async def test_create_subdirectory(self, tmp_path):
        result = await handle(
            action="create",
            path="sub/dir/test.docx",
            content_blocks=[{"type": "paragraph", "text": "nested"}],
        )
        assert "error" not in result
        assert (tmp_path / "sub" / "dir" / "test.docx").exists()


@_needs_docx
class TestRead:
    @pytest.mark.asyncio
    async def test_read_simple(self, tmp_path):
        await handle(
            action="create",
            path="read_me.docx",
            content_blocks=[
                {"type": "heading", "text": "Title", "level": 1},
                {"type": "paragraph", "text": "Body text"},
            ],
        )
        result = await handle(action="read", path="read_me.docx")
        assert "error" not in result
        assert "Title" in result["content"]
        assert "Body text" in result["content"]

    @pytest.mark.asyncio
    async def test_read_with_table(self, tmp_path):
        await handle(
            action="create",
            path="table.docx",
            content_blocks=[
                {"type": "table", "rows": [["Name", "Age"], ["Alice", "30"]]},
            ],
        )
        result = await handle(action="read", path="table.docx")
        assert "error" not in result
        assert "Alice" in result["content"]

    @pytest.mark.asyncio
    async def test_read_not_found(self):
        result = await handle(action="read", path="missing.docx")
        assert "error" in result
        assert "not found" in result["error"].lower()

    @pytest.mark.asyncio
    async def test_read_corrupt_file(self, tmp_path):
        corrupt = tmp_path / "corrupt.docx"
        corrupt.write_bytes(b"not a docx file")
        result = await handle(action="read", path="corrupt.docx")
        assert "error" in result


@_needs_docx
class TestEdit:
    @pytest.mark.asyncio
    async def test_edit_replace(self, tmp_path):
        await handle(
            action="create",
            path="edit_me.docx",
            content_blocks=[{"type": "paragraph", "text": "Hello world"}],
        )
        result = await handle(
            action="edit",
            path="edit_me.docx",
            replacements=[{"old": "Hello", "new": "Goodbye"}],
        )
        assert "error" not in result
        assert result["replacements_made"] >= 1

        read_result = await handle(action="read", path="edit_me.docx")
        assert "Goodbye" in read_result["content"]

    @pytest.mark.asyncio
    async def test_edit_append(self, tmp_path):
        await handle(
            action="create",
            path="append.docx",
            content_blocks=[{"type": "paragraph", "text": "Original"}],
        )
        result = await handle(
            action="edit",
            path="append.docx",
            content_blocks=[{"type": "paragraph", "text": "Appended"}],
        )
        assert "error" not in result
        assert result["blocks_appended"] == 1

    @pytest.mark.asyncio
    async def test_edit_not_found(self):
        result = await handle(
            action="edit",
            path="missing.docx",
            replacements=[{"old": "a", "new": "b"}],
        )
        assert "error" in result

    @pytest.mark.asyncio
    async def test_edit_no_operations(self, tmp_path):
        await handle(
            action="create",
            path="noop.docx",
            content_blocks=[{"type": "paragraph", "text": "text"}],
        )
        result = await handle(action="edit", path="noop.docx")
        assert "error" in result

    @pytest.mark.asyncio
    async def test_edit_too_many_blocks(self, tmp_path):
        await handle(
            action="create",
            path="big.docx",
            content_blocks=[{"type": "paragraph", "text": "text"}],
        )
        blocks = [{"type": "paragraph", "text": f"p{i}"} for i in range(_MAX_CONTENT_BLOCKS + 1)]
        result = await handle(action="edit", path="big.docx", content_blocks=blocks)
        assert "error" in result


@_needs_docx
class TestPathValidation:
    @pytest.mark.asyncio
    async def test_blocked_system_path_rejected(self):
        result = await handle(
            action="create",
            path="/etc/shadow",
            content_blocks=[{"type": "paragraph", "text": "x"}],
        )
        assert "error" in result

    @pytest.mark.asyncio
    async def test_null_bytes_rejected(self):
        result = await handle(action="read", path="test\x00.docx")
        assert "error" in result


@_needs_docx
class TestUnknownAction:
    @pytest.mark.asyncio
    async def test_unknown_action(self):
        result = await handle(action="delete", path="test.docx")
        assert "error" in result
        assert "Unknown action" in result["error"]


class TestGracefulDegradation:
    @pytest.mark.asyncio
    async def test_unavailable(self):
        with patch("anteroom.tools.office_docx.AVAILABLE", False):
            result = await handle(action="read", path="test.docx")
            assert "error" in result
            assert "pip install" in result["error"]


# ---------------------------------------------------------------------------
# New action tests (lib backend)
# ---------------------------------------------------------------------------


@_needs_docx
class TestHeadersFooters:
    async def _create_doc(self, name: str = "hf.docx") -> None:
        await handle(
            action="create",
            path=name,
            content_blocks=[{"type": "paragraph", "text": "Body text"}],
        )

    @pytest.mark.asyncio
    async def test_set_header_and_footer(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="headers_footers",
            path="hf.docx",
            operation="set",
            header_text="My Header",
            footer_text="My Footer",
        )
        assert "error" not in result
        assert "Set headers/footers" in result["result"]

    @pytest.mark.asyncio
    async def test_read_header_and_footer(self, tmp_path):
        await self._create_doc()
        await handle(
            action="headers_footers",
            path="hf.docx",
            operation="set",
            header_text="Test Header",
            footer_text="Test Footer",
        )
        result = await handle(
            action="headers_footers",
            path="hf.docx",
            operation="read",
        )
        assert "error" not in result
        assert "Test Header" in result["header"]
        assert "Test Footer" in result["footer"]
        assert result["section"] == 1

    @pytest.mark.asyncio
    async def test_read_default_empty(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="headers_footers",
            path="hf.docx",
            operation="read",
        )
        assert "error" not in result
        assert result["section"] == 1

    @pytest.mark.asyncio
    async def test_set_header_only(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="headers_footers",
            path="hf.docx",
            operation="set",
            header_text="Only Header",
        )
        assert "error" not in result

    @pytest.mark.asyncio
    async def test_set_footer_only(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="headers_footers",
            path="hf.docx",
            operation="set",
            footer_text="Only Footer",
        )
        assert "error" not in result

    @pytest.mark.asyncio
    async def test_invalid_section_index(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="headers_footers",
            path="hf.docx",
            operation="read",
            section_index=99,
        )
        assert "error" in result
        assert "Section 99 not found" in result["error"]

    @pytest.mark.asyncio
    async def test_unknown_operation(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="headers_footers",
            path="hf.docx",
            operation="delete",
        )
        assert "error" in result
        assert "Unknown operation" in result["error"]

    @pytest.mark.asyncio
    async def test_file_not_found(self):
        result = await handle(
            action="headers_footers",
            path="missing.docx",
            operation="read",
        )
        assert "error" in result
        assert "not found" in result["error"].lower()


@_needs_docx
class TestInsertImage:
    async def _create_doc(self, name: str = "img.docx") -> None:
        await handle(
            action="create",
            path=name,
            content_blocks=[{"type": "paragraph", "text": "Before image"}],
        )

    def _create_test_png(self, tmp_path, name: str = "test.png") -> str:
        """Create a minimal valid PNG file (1x1 red pixel)."""
        import struct
        import zlib

        def _chunk(chunk_type: bytes, data: bytes) -> bytes:
            c = chunk_type + data
            crc = struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)
            return struct.pack(">I", len(data)) + c + crc

        sig = b"\x89PNG\r\n\x1a\n"
        ihdr_data = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
        raw_row = b"\x00\xff\x00\x00"  # filter byte + RGB
        idat_data = zlib.compress(raw_row)

        png = sig + _chunk(b"IHDR", ihdr_data) + _chunk(b"IDAT", idat_data) + _chunk(b"IEND", b"")
        path = tmp_path / name
        path.write_bytes(png)
        return name

    @pytest.mark.asyncio
    async def test_insert_image_basic(self, tmp_path):
        await self._create_doc()
        img_name = self._create_test_png(tmp_path)
        result = await handle(
            action="insert_image",
            path="img.docx",
            image_path=img_name,
        )
        assert "error" not in result
        assert "Inserted image" in result["result"]
        assert result["path"] == "img.docx"

    @pytest.mark.asyncio
    async def test_insert_image_with_width(self, tmp_path):
        await self._create_doc()
        img_name = self._create_test_png(tmp_path)
        result = await handle(
            action="insert_image",
            path="img.docx",
            image_path=img_name,
            width_inches=3.0,
        )
        assert "error" not in result

    @pytest.mark.asyncio
    async def test_insert_image_missing_path(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="insert_image",
            path="img.docx",
        )
        assert "error" in result
        assert "image_path is required" in result["error"]

    @pytest.mark.asyncio
    async def test_insert_image_not_found(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="insert_image",
            path="img.docx",
            image_path="nonexistent.png",
        )
        assert "error" in result
        assert "not found" in result["error"].lower()

    @pytest.mark.asyncio
    async def test_insert_image_doc_not_found(self):
        result = await handle(
            action="insert_image",
            path="missing.docx",
            image_path="test.png",
        )
        assert "error" in result
        assert "not found" in result["error"].lower()


@_needs_docx
class TestStyles:
    async def _create_doc(self, name: str = "styles.docx") -> None:
        await handle(
            action="create",
            path=name,
            content_blocks=[
                {"type": "heading", "text": "Title", "level": 1},
                {"type": "paragraph", "text": "Normal paragraph"},
            ],
        )

    @pytest.mark.asyncio
    async def test_list_styles(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="styles",
            path="styles.docx",
            operation="list",
        )
        assert "error" not in result
        assert "styles" in result
        assert result["count"] > 0
        style_names = [s["name"] for s in result["styles"]]
        assert any("Normal" in n for n in style_names)

    @pytest.mark.asyncio
    async def test_list_styles_default_operation(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="styles",
            path="styles.docx",
        )
        assert "error" not in result
        assert "styles" in result

    @pytest.mark.asyncio
    async def test_read_style(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="styles",
            path="styles.docx",
            operation="read",
            style_name="Normal",
        )
        assert "error" not in result
        assert result["style"]["name"] == "Normal"

    @pytest.mark.asyncio
    async def test_read_style_not_found(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="styles",
            path="styles.docx",
            operation="read",
            style_name="NonexistentStyleXYZ",
        )
        assert "error" in result
        assert "not found" in result["error"]

    @pytest.mark.asyncio
    async def test_read_style_missing_name(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="styles",
            path="styles.docx",
            operation="read",
        )
        assert "error" in result
        assert "style_name required" in result["error"]

    @pytest.mark.asyncio
    async def test_set_style(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="styles",
            path="styles.docx",
            operation="set",
            style_name="Heading 2",
            paragraph_index=1,
        )
        assert "error" not in result
        assert "Applied style" in result["result"]

    @pytest.mark.asyncio
    async def test_set_style_out_of_range(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="styles",
            path="styles.docx",
            operation="set",
            style_name="Normal",
            paragraph_index=999,
        )
        assert "error" in result
        assert "out of range" in result["error"]

    @pytest.mark.asyncio
    async def test_set_style_missing_params(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="styles",
            path="styles.docx",
            operation="set",
        )
        assert "error" in result
        assert "style_name and paragraph_index required" in result["error"]

    @pytest.mark.asyncio
    async def test_unknown_operation(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="styles",
            path="styles.docx",
            operation="delete",
        )
        assert "error" in result
        assert "Unknown operation" in result["error"]

    @pytest.mark.asyncio
    async def test_file_not_found(self):
        result = await handle(
            action="styles",
            path="missing.docx",
            operation="list",
        )
        assert "error" in result


@_needs_docx
class TestPageSetup:
    async def _create_doc(self, name: str = "page.docx") -> None:
        await handle(
            action="create",
            path=name,
            content_blocks=[{"type": "paragraph", "text": "Content"}],
        )

    @pytest.mark.asyncio
    async def test_read_page_setup(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="page_setup",
            path="page.docx",
            operation="read",
        )
        assert "error" not in result
        assert result["result"] == "Read page setup"
        assert "orientation" in result
        assert "top_margin" in result
        assert "bottom_margin" in result
        assert "left_margin" in result
        assert "right_margin" in result
        assert "page_width" in result
        assert "page_height" in result

    @pytest.mark.asyncio
    async def test_set_orientation_landscape(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="page_setup",
            path="page.docx",
            orientation="landscape",
        )
        assert "error" not in result
        assert "Updated page setup" in result["result"]

        read_result = await handle(
            action="page_setup",
            path="page.docx",
            operation="read",
        )
        assert read_result["orientation"] == "landscape"

    @pytest.mark.asyncio
    async def test_set_orientation_portrait(self, tmp_path):
        await self._create_doc()
        # First set to landscape, then back to portrait
        await handle(
            action="page_setup",
            path="page.docx",
            orientation="landscape",
        )
        result = await handle(
            action="page_setup",
            path="page.docx",
            orientation="portrait",
        )
        assert "error" not in result
        read_result = await handle(
            action="page_setup",
            path="page.docx",
            operation="read",
        )
        assert read_result["orientation"] == "portrait"

    @pytest.mark.asyncio
    async def test_set_margins(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="page_setup",
            path="page.docx",
            margins={"top": 0.5, "bottom": 0.5, "left": 0.75, "right": 0.75},
        )
        assert "error" not in result
        assert "Updated page setup" in result["result"]

    @pytest.mark.asyncio
    async def test_set_paper_size(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="page_setup",
            path="page.docx",
            paper_size="a4",
        )
        assert "error" not in result

    @pytest.mark.asyncio
    async def test_file_not_found(self):
        result = await handle(
            action="page_setup",
            path="missing.docx",
            operation="read",
        )
        assert "error" in result


@_needs_docx
class TestSections:
    async def _create_doc(self, name: str = "sections.docx") -> None:
        await handle(
            action="create",
            path=name,
            content_blocks=[{"type": "paragraph", "text": "Section 1 content"}],
        )

    @pytest.mark.asyncio
    async def test_list_sections(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="sections",
            path="sections.docx",
            operation="list",
        )
        assert "error" not in result
        assert result["count"] >= 1
        assert "sections" in result
        section = result["sections"][0]
        assert "index" in section
        assert "start_type" in section
        assert "page_width" in section
        assert "page_height" in section

    @pytest.mark.asyncio
    async def test_list_sections_default_operation(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="sections",
            path="sections.docx",
        )
        assert "error" not in result
        assert "sections" in result

    @pytest.mark.asyncio
    async def test_add_section(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="sections",
            path="sections.docx",
            operation="add",
            start_type="new_page",
        )
        assert "error" not in result
        assert "Added new_page section break" in result["result"]
        assert result["total_sections"] == 2

    @pytest.mark.asyncio
    async def test_add_continuous_section(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="sections",
            path="sections.docx",
            operation="add",
            start_type="continuous",
        )
        assert "error" not in result
        assert "Added continuous section break" in result["result"]

    @pytest.mark.asyncio
    async def test_add_section_default_type(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="sections",
            path="sections.docx",
            operation="add",
        )
        assert "error" not in result
        assert "new_page" in result["result"]

    @pytest.mark.asyncio
    async def test_unknown_operation(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="sections",
            path="sections.docx",
            operation="delete",
        )
        assert "error" in result
        assert "Unknown operation" in result["error"]

    @pytest.mark.asyncio
    async def test_file_not_found(self):
        result = await handle(
            action="sections",
            path="missing.docx",
            operation="list",
        )
        assert "error" in result


@_needs_docx
class TestFindRegex:
    async def _create_doc(self, name: str = "regex.docx") -> None:
        await handle(
            action="create",
            path=name,
            content_blocks=[
                {"type": "paragraph", "text": "Hello world"},
                {"type": "paragraph", "text": "Hello again"},
                {"type": "paragraph", "text": "Goodbye world"},
            ],
        )

    @pytest.mark.asyncio
    async def test_find_simple_pattern(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="find_regex",
            path="regex.docx",
            pattern="Hello",
        )
        assert "error" not in result
        assert result["count"] == 2
        assert len(result["matches"]) == 2
        assert result["truncated"] is False

    @pytest.mark.asyncio
    async def test_find_regex_pattern(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="find_regex",
            path="regex.docx",
            pattern=r"Hello\s\w+",
        )
        assert "error" not in result
        assert result["count"] == 2
        for match in result["matches"]:
            assert "text" in match
            assert "paragraph" in match
            assert "start" in match
            assert "end" in match

    @pytest.mark.asyncio
    async def test_find_no_matches(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="find_regex",
            path="regex.docx",
            pattern="ZZZZZ",
        )
        assert "error" not in result
        assert result["count"] == 0
        assert result["matches"] == []

    @pytest.mark.asyncio
    async def test_find_and_replace(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="find_regex",
            path="regex.docx",
            pattern="Hello",
            replace_with="Hi",
        )
        assert "error" not in result
        assert result["replacements"] >= 2
        assert result["path"] == "regex.docx"

        read_result = await handle(action="read", path="regex.docx")
        assert "Hi" in read_result["content"]
        assert "Hello" not in read_result["content"]

    @pytest.mark.asyncio
    async def test_find_replace_regex(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="find_regex",
            path="regex.docx",
            pattern=r"Hello\s(\w+)",
            replace_with=r"Greetings \1",
        )
        assert "error" not in result
        assert result["replacements"] >= 1

    @pytest.mark.asyncio
    async def test_missing_pattern(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="find_regex",
            path="regex.docx",
        )
        assert "error" in result
        assert "pattern is required" in result["error"]

    @pytest.mark.asyncio
    async def test_invalid_regex(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="find_regex",
            path="regex.docx",
            pattern="[invalid",
        )
        assert "error" in result
        assert "Invalid regex" in result["error"]

    @pytest.mark.asyncio
    async def test_pattern_too_long(self, tmp_path):
        await self._create_doc()
        result = await handle(
            action="find_regex",
            path="regex.docx",
            pattern="a" * 201,
        )
        assert "error" in result
        assert "too long" in result["error"]

    @pytest.mark.asyncio
    async def test_file_not_found(self):
        result = await handle(
            action="find_regex",
            path="missing.docx",
            pattern="test",
        )
        assert "error" in result


# ---------------------------------------------------------------------------
# COM-only action tests (verify they return the COM-only error on lib backend)
# ---------------------------------------------------------------------------


@_needs_docx
class TestTrackChangesCOMOnly:
    @pytest.mark.asyncio
    async def test_returns_com_only_error(self, tmp_path):
        await handle(
            action="create",
            path="tc.docx",
            content_blocks=[{"type": "paragraph", "text": "text"}],
        )
        result = await handle(action="track_changes", path="tc.docx")
        assert "error" in result
        assert "COM backend" in result["error"]


@_needs_docx
class TestCommentsCOMOnly:
    @pytest.mark.asyncio
    async def test_returns_com_only_error(self, tmp_path):
        await handle(
            action="create",
            path="comments.docx",
            content_blocks=[{"type": "paragraph", "text": "text"}],
        )
        result = await handle(action="comments", path="comments.docx")
        assert "error" in result
        assert "COM backend" in result["error"]


@_needs_docx
class TestBookmarksCOMOnly:
    @pytest.mark.asyncio
    async def test_returns_com_only_error(self, tmp_path):
        await handle(
            action="create",
            path="bm.docx",
            content_blocks=[{"type": "paragraph", "text": "text"}],
        )
        result = await handle(action="bookmarks", path="bm.docx")
        assert "error" in result
        assert "COM backend" in result["error"]


@_needs_docx
class TestTocCOMOnly:
    @pytest.mark.asyncio
    async def test_returns_com_only_error(self, tmp_path):
        await handle(
            action="create",
            path="toc.docx",
            content_blocks=[{"type": "paragraph", "text": "text"}],
        )
        result = await handle(action="toc", path="toc.docx")
        assert "error" in result
        assert "COM backend" in result["error"]


@_needs_docx
class TestExportPdfCOMOnly:
    @pytest.mark.asyncio
    async def test_returns_com_only_error(self, tmp_path):
        await handle(
            action="create",
            path="export.docx",
            content_blocks=[{"type": "paragraph", "text": "text"}],
        )
        result = await handle(action="export_pdf", path="export.docx")
        assert "error" in result
        assert "COM backend" in result["error"]


class TestComDispatchErrorHandling:
    @pytest.mark.asyncio
    async def test_dispatch_com_returns_error_dict_on_exception(self):
        from unittest.mock import AsyncMock, MagicMock

        mock_manager = MagicMock()
        mock_manager.run_com = AsyncMock(side_effect=RuntimeError("Access denied by security policy"))
        mock_com_mod = MagicMock()
        mock_com_mod.get_manager.return_value = mock_manager
        mock_com_mod.COM_AVAILABLE = True

        with (
            patch("anteroom.tools.office_docx.AVAILABLE", True),
            patch("anteroom.tools.office_docx._BACKEND", "com"),
            patch("anteroom.tools.office_docx._com_mod", mock_com_mod),
        ):
            result = await handle(action="edit", path="test.docx", replacements=[{"old": "a", "new": "b"}])

        assert "error" in result
        assert "Access denied by security policy" in result["error"]
        assert "RuntimeError" in result["error"]


# ---------------------------------------------------------------------------
# Enriched read output tests
# ---------------------------------------------------------------------------


@_needs_docx
class TestEnrichedRead:
    @pytest.mark.asyncio
    async def test_read_includes_paragraph_formatting(self, tmp_path):
        import docx as _docx
        from docx.shared import Pt

        doc = _docx.Document()
        para = doc.add_paragraph()
        run = para.add_run("Bold text")
        run.bold = True
        run.font.size = Pt(14)
        doc.save(str(tmp_path / "fmt.docx"))
        result = await handle(action="read", path="fmt.docx")
        assert "error" not in result
        assert "bold" in result["content"]
        assert "size=14" in result["content"]
        assert "Bold text" in result["content"]

    @pytest.mark.asyncio
    async def test_read_includes_italic(self, tmp_path):
        import docx as _docx

        doc = _docx.Document()
        para = doc.add_paragraph()
        run = para.add_run("Italic text")
        run.italic = True
        doc.save(str(tmp_path / "italic.docx"))
        result = await handle(action="read", path="italic.docx")
        assert "italic" in result["content"]

    @pytest.mark.asyncio
    async def test_read_includes_underline(self, tmp_path):
        import docx as _docx

        doc = _docx.Document()
        para = doc.add_paragraph()
        run = para.add_run("Underlined")
        run.underline = True
        doc.save(str(tmp_path / "ul.docx"))
        result = await handle(action="read", path="ul.docx")
        assert "underline" in result["content"]

    @pytest.mark.asyncio
    async def test_read_includes_style_name(self, tmp_path):
        await handle(
            action="create",
            path="style.docx",
            content_blocks=[
                {"type": "heading", "text": "Title", "level": 1},
                {"type": "paragraph", "text": "Normal paragraph"},
            ],
        )
        result = await handle(action="read", path="style.docx")
        assert "style=Heading 1" in result["content"]

    @pytest.mark.asyncio
    async def test_read_includes_font_name(self, tmp_path):
        import docx as _docx

        doc = _docx.Document()
        para = doc.add_paragraph()
        run = para.add_run("Courier text")
        run.font.name = "Courier New"
        doc.save(str(tmp_path / "font.docx"))
        result = await handle(action="read", path="font.docx")
        assert "font=Courier New" in result["content"]

    @pytest.mark.asyncio
    async def test_read_table_cell_multi_paragraph(self, tmp_path):
        import docx as _docx

        doc = _docx.Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        cell.paragraphs[0].text = "First paragraph"
        cell.add_paragraph("Second paragraph")
        doc.save(str(tmp_path / "tbl.docx"))
        result = await handle(action="read", path="tbl.docx")
        assert "error" not in result
        assert "P0" in result["content"]
        assert "P1" in result["content"]
        assert "First paragraph" in result["content"]
        assert "Second paragraph" in result["content"]

    @pytest.mark.asyncio
    async def test_read_table_dimensions(self, tmp_path):
        await handle(
            action="create",
            path="tbl_dim.docx",
            content_blocks=[
                {"type": "table", "rows": [["A", "B", "C"], ["1", "2", "3"]]},
            ],
        )
        result = await handle(action="read", path="tbl_dim.docx")
        assert "2 rows x 3 cols" in result["content"]

    @pytest.mark.asyncio
    async def test_read_returns_section_count(self, tmp_path):
        await handle(
            action="create",
            path="sec.docx",
            content_blocks=[{"type": "paragraph", "text": "content"}],
        )
        result = await handle(action="read", path="sec.docx")
        assert "error" not in result
        assert "sections" in result
        assert result["sections"] >= 1

    @pytest.mark.asyncio
    async def test_read_returns_sections_info(self, tmp_path):
        await handle(
            action="create",
            path="secinfo.docx",
            content_blocks=[{"type": "paragraph", "text": "content"}],
        )
        result = await handle(action="read", path="secinfo.docx")
        assert "error" not in result
        assert "sections_info" in result
        si = result["sections_info"][0]
        assert "orientation" in si
        assert "page_width" in si
        assert "page_height" in si

    @pytest.mark.asyncio
    async def test_read_returns_doc_properties(self, tmp_path):
        import docx as _docx

        doc = _docx.Document()
        doc.core_properties.title = "Test Title"
        doc.core_properties.author = "Test Author"
        doc.add_paragraph("content")
        doc.save(str(tmp_path / "props.docx"))
        result = await handle(action="read", path="props.docx")
        assert "error" not in result
        assert "properties" in result
        assert result["properties"]["title"] == "Test Title"
        assert result["properties"]["author"] == "Test Author"

    @pytest.mark.asyncio
    async def test_read_headers_footers_shown(self, tmp_path):
        await handle(
            action="create",
            path="hf_read.docx",
            content_blocks=[{"type": "paragraph", "text": "Body"}],
        )
        await handle(
            action="headers_footers",
            path="hf_read.docx",
            operation="set",
            header_text="Report Header",
            footer_text="Page Footer",
        )
        result = await handle(action="read", path="hf_read.docx")
        assert "error" not in result
        assert "Report Header" in result["content"]
        assert "Page Footer" in result["content"]

    @pytest.mark.asyncio
    async def test_read_paragraph_index_format(self, tmp_path):
        """Verify that paragraphs use P0, P1, P2... indexing."""
        await handle(
            action="create",
            path="pidx.docx",
            content_blocks=[
                {"type": "paragraph", "text": "First"},
                {"type": "paragraph", "text": "Second"},
                {"type": "paragraph", "text": "Third"},
            ],
        )
        result = await handle(action="read", path="pidx.docx")
        assert "P0" in result["content"] or "P1" in result["content"]


# ---------------------------------------------------------------------------
# template_fill tests
# ---------------------------------------------------------------------------


@_needs_docx
class TestTemplateFill:
    @pytest.mark.asyncio
    async def test_template_fill_basic(self, tmp_path):
        await handle(
            action="create",
            path="tpl.docx",
            content_blocks=[
                {"type": "paragraph", "text": "Hello {{name}}, welcome to {{company}}!"},
            ],
        )
        result = await handle(
            action="template_fill",
            path="tpl.docx",
            template_data={"name": "Alice", "company": "Acme Corp"},
        )
        assert "error" not in result
        assert result["tokens_replaced"] >= 2
        assert result["keys_processed"] == 2
        read = await handle(action="read", path="tpl.docx")
        assert "Alice" in read["content"]
        assert "Acme Corp" in read["content"]
        assert "{{name}}" not in read["content"]

    @pytest.mark.asyncio
    async def test_template_fill_in_tables(self, tmp_path):
        await handle(
            action="create",
            path="tpl_tbl.docx",
            content_blocks=[
                {"type": "table", "rows": [["Name", "Value"], ["{{item}}", "{{price}}"]]},
            ],
        )
        result = await handle(
            action="template_fill",
            path="tpl_tbl.docx",
            template_data={"item": "Widget", "price": "$9.99"},
        )
        assert "error" not in result
        assert result["tokens_replaced"] >= 2
        read = await handle(action="read", path="tpl_tbl.docx")
        assert "Widget" in read["content"]
        assert "$9.99" in read["content"]

    @pytest.mark.asyncio
    async def test_template_fill_in_headers_footers(self, tmp_path):
        await handle(
            action="create",
            path="tpl_hf.docx",
            content_blocks=[{"type": "paragraph", "text": "Body"}],
        )
        await handle(
            action="headers_footers",
            path="tpl_hf.docx",
            operation="set",
            header_text="Report: {{title}}",
            footer_text="Page {{page}}",
        )
        result = await handle(
            action="template_fill",
            path="tpl_hf.docx",
            template_data={"title": "Q4 Report", "page": "1"},
        )
        assert "error" not in result
        assert result["tokens_replaced"] >= 2

    @pytest.mark.asyncio
    async def test_template_fill_no_data(self):
        result = await handle(action="template_fill", path="test.docx")
        assert "error" in result
        assert "template_data is required" in result["error"]

    @pytest.mark.asyncio
    async def test_template_fill_empty_data(self):
        result = await handle(action="template_fill", path="test.docx", template_data={})
        assert "error" in result

    @pytest.mark.asyncio
    async def test_template_fill_too_many_keys(self):
        big_data = {f"key{i}": f"val{i}" for i in range(_MAX_EDIT_OPS + 1)}
        result = await handle(action="template_fill", path="test.docx", template_data=big_data)
        assert "error" in result
        assert "Too many" in result["error"]

    @pytest.mark.asyncio
    async def test_template_fill_file_not_found(self):
        result = await handle(
            action="template_fill",
            path="missing.docx",
            template_data={"key": "val"},
        )
        assert "error" in result
        assert "not found" in result["error"].lower()

    @pytest.mark.asyncio
    async def test_template_fill_no_matches(self, tmp_path):
        await handle(
            action="create",
            path="nomatch.docx",
            content_blocks=[{"type": "paragraph", "text": "no tokens here"}],
        )
        result = await handle(
            action="template_fill",
            path="nomatch.docx",
            template_data={"missing": "value"},
        )
        assert "error" not in result
        assert result["tokens_replaced"] == 0

    @pytest.mark.asyncio
    async def test_template_fill_preserves_formatting(self, tmp_path):
        import docx as _docx
        from docx.shared import Pt

        doc = _docx.Document()
        para = doc.add_paragraph()
        run = para.add_run("{{name}}")
        run.bold = True
        run.font.size = Pt(16)
        doc.save(str(tmp_path / "fmt_tpl.docx"))

        result = await handle(
            action="template_fill",
            path="fmt_tpl.docx",
            template_data={"name": "Alice"},
        )
        assert "error" not in result
        assert result["tokens_replaced"] >= 1

        # Verify formatting preserved
        doc2 = _docx.Document(str(tmp_path / "fmt_tpl.docx"))
        assert doc2.paragraphs[0].runs[0].bold is True
        assert doc2.paragraphs[0].runs[0].font.size == Pt(16)
        assert doc2.paragraphs[0].runs[0].text == "Alice"

    @pytest.mark.asyncio
    async def test_template_fill_multiple_occurrences(self, tmp_path):
        await handle(
            action="create",
            path="multi.docx",
            content_blocks=[
                {"type": "paragraph", "text": "{{name}} likes {{name}}"},
            ],
        )
        result = await handle(
            action="template_fill",
            path="multi.docx",
            template_data={"name": "Alice"},
        )
        assert "error" not in result
        assert result["tokens_replaced"] >= 1
        read = await handle(action="read", path="multi.docx")
        assert "{{name}}" not in read["content"]
        assert "Alice" in read["content"]


# ---------------------------------------------------------------------------
# New action definition tests
# ---------------------------------------------------------------------------


class TestNewActionDefinitions:
    def test_template_fill_in_definition(self):
        actions = DEFINITION["parameters"]["properties"]["action"]["enum"]
        assert "template_fill" in actions

    def test_template_data_param(self):
        props = DEFINITION["parameters"]["properties"]
        assert "template_data" in props
        assert props["template_data"]["type"] == "object"

    def test_description_mentions_template_fill(self):
        assert "template fill" in DEFINITION["description"].lower()

    def test_description_mentions_formatting_annotations(self):
        assert "formatting annotations" in DEFINITION["description"].lower()
