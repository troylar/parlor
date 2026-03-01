"""DOCX (Word document) create/read/edit tool.

Backends:
- COM (Windows + Office + pywin32): full Office object model
- Library (python-docx): cross-platform XML manipulation

Install: ``pip install anteroom[office]`` or ``pip install anteroom[office-com]``
"""

from __future__ import annotations

import json
import os
import re
import sys
from typing import Any

from .security import validate_path

_BACKEND: str | None = None
_com_mod: Any = None

if sys.platform == "win32":
    try:
        from . import office_com as _com_mod_import

        if _com_mod_import.COM_AVAILABLE:
            _com_mod = _com_mod_import
            _BACKEND = "com"
    except ImportError:
        pass

if _BACKEND is None:
    try:
        import docx  # noqa: F401

        _BACKEND = "lib"
    except ImportError:
        pass

AVAILABLE = _BACKEND is not None

_MAX_OUTPUT = 100_000
_MAX_CONTENT_BLOCKS = 200

_ALL_ACTIONS = [
    "create",
    "read",
    "edit",
    "track_changes",
    "comments",
    "headers_footers",
    "insert_image",
    "styles",
    "export_pdf",
    "page_setup",
    "sections",
    "bookmarks",
    "toc",
    "find_regex",
]

DEFINITION: dict[str, Any] = {
    "name": "docx",
    "description": (
        "Create, read, edit, and manipulate Word documents (.docx). "
        "Supports track changes, comments, headers/footers, images, styles, "
        "PDF export, page setup, sections, bookmarks, table of contents, and regex find."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "action": {
                "type": "string",
                "enum": _ALL_ACTIONS,
                "description": "Operation to perform.",
            },
            "path": {
                "type": "string",
                "description": "File path (relative to working directory or absolute).",
            },
            "content_blocks": {
                "type": "array",
                "description": (
                    "List of content blocks for create/edit append. "
                    "Each block: {type: 'heading'|'paragraph'|'table', text?: str, "
                    "level?: int, rows?: [[str]]}."
                ),
                "items": {"type": "object"},
            },
            "replacements": {
                "type": "array",
                "description": "List of {old: str, new: str} for find/replace in edit action.",
                "items": {"type": "object"},
            },
            "operation": {
                "type": "string",
                "enum": [
                    "accept",
                    "reject",
                    "list",
                    "add",
                    "read",
                    "delete",
                    "set",
                    "enable",
                    "disable",
                ],
                "description": "Sub-operation for actions that support multiple modes.",
            },
            "comment_text": {
                "type": "string",
                "description": "Comment text for comments action.",
            },
            "author": {
                "type": "string",
                "description": "Comment author name.",
            },
            "bookmark_name": {
                "type": "string",
                "description": "Bookmark name for bookmarks and insert_image actions.",
            },
            "section_index": {
                "type": "integer",
                "description": "Section number (1-based) for headers_footers and sections actions.",
            },
            "header_text": {
                "type": "string",
                "description": "Header text for headers_footers action.",
            },
            "footer_text": {
                "type": "string",
                "description": "Footer text for headers_footers action.",
            },
            "image_path": {
                "type": "string",
                "description": "Image file path for insert_image action.",
            },
            "width_inches": {
                "type": "number",
                "description": "Image width in inches for insert_image action.",
            },
            "style_name": {
                "type": "string",
                "description": "Style name for styles action.",
            },
            "paragraph_index": {
                "type": "integer",
                "description": "Paragraph index (0-based) for styles apply operation.",
            },
            "output_path": {
                "type": "string",
                "description": "Output file path for export_pdf action.",
            },
            "orientation": {
                "type": "string",
                "enum": ["portrait", "landscape"],
                "description": "Page orientation for page_setup action.",
            },
            "margins": {
                "type": "object",
                "description": (
                    "Page margins in inches for page_setup: {top?: float, bottom?: float, left?: float, right?: float}."
                ),
            },
            "paper_size": {
                "type": "string",
                "enum": ["letter", "a4", "legal"],
                "description": "Paper size for page_setup action.",
            },
            "start_type": {
                "type": "string",
                "enum": ["new_page", "continuous", "even_page", "odd_page"],
                "description": "Section start type for sections add operation.",
            },
            "pattern": {
                "type": "string",
                "description": "Regex pattern for find_regex action.",
            },
            "replace_with": {
                "type": "string",
                "description": "Replacement text for find_regex action.",
            },
        },
        "required": ["action", "path"],
    },
}


def _open_document_lib(
    resolved: str,
    display_path: str,
) -> tuple[Any, str | None]:
    """Open a document with python-docx, returning (doc, error)."""
    import docx as _docx

    if not os.path.isfile(resolved):
        return None, f"File not found: {display_path}"
    try:
        document = _docx.Document(resolved)
        return document, None
    except Exception as exc:
        return None, f"Unable to read DOCX file: {display_path} ({exc})"


def _open_document_com(
    manager: Any,
    resolved: str,
    display_path: str,
    read_only: bool = False,
) -> tuple[Any, Any, str | None]:
    """Open document via COM. Returns (word, doc, error)."""
    if not os.path.isfile(resolved):
        return None, None, f"File not found: {display_path}"
    word = manager.get_app("Word.Application")
    try:
        doc = word.Documents.Open(os.path.abspath(resolved), ReadOnly=read_only)
        return word, doc, None
    except Exception as exc:
        return None, None, f"Unable to read DOCX file: {display_path} ({exc})"


async def handle(action: str, path: str, **kwargs: Any) -> dict[str, Any]:
    if not AVAILABLE:
        return {"error": "No docx backend available. Install with: pip install anteroom[office]"}

    working_dir = kwargs.pop("_working_dir", None) or os.getcwd()
    resolved, error = validate_path(path, working_dir)
    if error:
        return {"error": error}

    if _BACKEND == "com":
        return await _dispatch_com(action, resolved, path, working_dir=working_dir, **kwargs)

    _lib_dispatch: dict[str, Any] = {
        "create": _create_lib,
        "read": _read_lib,
        "edit": _edit_lib,
        "track_changes": _track_changes_lib,
        "comments": _comments_lib,
        "headers_footers": _headers_footers_lib,
        "insert_image": _insert_image_lib,
        "styles": _styles_lib,
        "export_pdf": _export_pdf_lib,
        "page_setup": _page_setup_lib,
        "sections": _sections_lib,
        "bookmarks": _bookmarks_lib,
        "toc": _toc_lib,
        "find_regex": _find_regex_lib,
    }

    handler = _lib_dispatch.get(action)
    if handler is None:
        return {"error": f"Unknown action: {action}. Available: {', '.join(_ALL_ACTIONS)}"}
    return dict(handler(resolved, path, working_dir=working_dir, **kwargs))


# ---------------------------------------------------------------------------
# COM backend
# ---------------------------------------------------------------------------


async def _dispatch_com(
    action: str,
    resolved: str,
    display_path: str,
    *,
    working_dir: str,
    **kwargs: Any,
) -> dict[str, Any]:
    manager = _com_mod.get_manager()

    _com_dispatch: dict[str, Any] = {
        "create": _create_com,
        "read": _read_com,
        "edit": _edit_com,
        "track_changes": _track_changes_com,
        "comments": _comments_com,
        "headers_footers": _headers_footers_com,
        "insert_image": _insert_image_com,
        "styles": _styles_com,
        "export_pdf": _export_pdf_com,
        "page_setup": _page_setup_com,
        "sections": _sections_com,
        "bookmarks": _bookmarks_com,
        "toc": _toc_com,
        "find_regex": _find_regex_com,
    }

    handler = _com_dispatch.get(action)
    if handler is None:
        return {"error": f"Unknown action: {action}. Available: {', '.join(_ALL_ACTIONS)}"}
    try:
        return dict(await manager.run_com(handler, manager, resolved, display_path, working_dir=working_dir, **kwargs))
    except Exception as exc:
        return {"error": f"COM {action} failed on {display_path}: {type(exc).__name__}: {exc}"}


# --- create ---


def _create_com(manager: Any, resolved: str, display_path: str, **kwargs: Any) -> dict[str, Any]:
    content_blocks: list[dict[str, Any]] = kwargs.get("content_blocks") or []
    if not content_blocks:
        return {"error": "content_blocks is required for create action"}
    if len(content_blocks) > _MAX_CONTENT_BLOCKS:
        return {"error": f"Too many content blocks (max {_MAX_CONTENT_BLOCKS})"}

    word = manager.get_app("Word.Application")
    doc = word.Documents.Add()

    try:
        rng = doc.Content
        for block in content_blocks:
            block_type = block.get("type", "paragraph")
            text = block.get("text", "")

            if block_type == "heading":
                level = block.get("level", 1)
                level = max(0, min(9, level))
                rng.InsertAfter(text + "\n")
                rng.Start = rng.End - len(text) - 1
                if level > 0:
                    rng.Style = f"Heading {level}"
                else:
                    rng.Style = "Title"
                rng.Start = rng.End

            elif block_type == "table":
                rows = block.get("rows", [])
                if not rows:
                    continue
                cols = max(len(r) for r in rows) if rows else 1
                rng.Start = rng.End
                table = doc.Tables.Add(rng, len(rows), cols)
                for i, row_data in enumerate(rows):
                    for j, cell_text in enumerate(row_data):
                        if j < cols:
                            table.Cell(i + 1, j + 1).Range.Text = str(cell_text)
                rng.Start = doc.Content.End - 1
                rng.End = rng.Start

            else:
                rng.InsertAfter(text + "\n")
                rng.Start = rng.End

        os.makedirs(os.path.dirname(resolved) or ".", exist_ok=True)
        doc.SaveAs2(os.path.abspath(resolved))
    finally:
        doc.Close(SaveChanges=False)

    return {"result": f"Created {display_path}", "path": display_path, "blocks_written": len(content_blocks)}


def _create_lib(resolved: str, display_path: str, **kwargs: Any) -> dict[str, Any]:
    import docx as _docx

    content_blocks: list[dict[str, Any]] = kwargs.get("content_blocks") or []
    if not content_blocks:
        return {"error": "content_blocks is required for create action"}
    if len(content_blocks) > _MAX_CONTENT_BLOCKS:
        return {"error": f"Too many content blocks (max {_MAX_CONTENT_BLOCKS})"}

    document = _docx.Document()

    for block in content_blocks:
        block_type = block.get("type", "paragraph")
        if block_type == "heading":
            level = block.get("level", 1)
            level = max(0, min(9, level))
            document.add_heading(block.get("text", ""), level=level)
        elif block_type == "paragraph":
            document.add_paragraph(block.get("text", ""))
        elif block_type == "table":
            rows = block.get("rows", [])
            if not rows:
                continue
            cols = max(len(r) for r in rows) if rows else 1
            table = document.add_table(rows=len(rows), cols=cols)
            for i, row_data in enumerate(rows):
                for j, cell_text in enumerate(row_data):
                    if j < cols:
                        table.rows[i].cells[j].text = str(cell_text)
        else:
            document.add_paragraph(block.get("text", ""))

    os.makedirs(os.path.dirname(resolved) or ".", exist_ok=True)
    document.save(resolved)
    return {"result": f"Created {display_path}", "path": display_path, "blocks_written": len(content_blocks)}


# --- read ---


def _read_com(manager: Any, resolved: str, display_path: str, **kwargs: Any) -> dict[str, Any]:
    if not os.path.isfile(resolved):
        return {"error": f"File not found: {display_path}"}

    word = manager.get_app("Word.Application")
    try:
        doc = word.Documents.Open(os.path.abspath(resolved), ReadOnly=True)
    except Exception as exc:
        return {"error": f"Unable to read DOCX file: {display_path} ({type(exc).__name__}: {exc})"}

    try:
        output_parts: list[str] = []

        for para in doc.Paragraphs:
            style_name = para.Style.NameLocal
            text = para.Range.Text.rstrip("\r\x07")
            if style_name.startswith("Heading"):
                try:
                    level = int(style_name.split()[-1])
                except (ValueError, IndexError):
                    level = 1
                output_parts.append(f"{'#' * level} {text}")
            elif text.strip():
                output_parts.append(text)

        for i in range(1, doc.Tables.Count + 1):
            output_parts.append(f"\n[Table {i}]")
            table = doc.Tables(i)
            table_rows: list[list[str]] = []
            for r in range(1, table.Rows.Count + 1):
                row_data: list[str] = []
                for c in range(1, table.Columns.Count + 1):
                    cell_text = table.Cell(r, c).Range.Text.rstrip("\r\x07")
                    row_data.append(cell_text)
                table_rows.append(row_data)
            output_parts.append(json.dumps(table_rows, ensure_ascii=False))

        content = "\n".join(output_parts)
        if len(content) > _MAX_OUTPUT:
            content = content[:_MAX_OUTPUT] + "\n... (truncated)"

        para_count = doc.Paragraphs.Count
        table_count = doc.Tables.Count
    finally:
        doc.Close(SaveChanges=False)

    return {
        "content": content,
        "paragraphs": para_count,
        "tables": table_count,
    }


def _read_lib(resolved: str, display_path: str, **kwargs: Any) -> dict[str, Any]:
    import docx as _docx

    if not os.path.isfile(resolved):
        return {"error": f"File not found: {display_path}"}

    try:
        document = _docx.Document(resolved)
    except Exception as exc:
        return {"error": f"Unable to read DOCX file: {display_path} ({type(exc).__name__}: {exc})"}

    output_parts: list[str] = []

    for para in document.paragraphs:
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.split()[-1])
            except (ValueError, IndexError):
                level = 1
            output_parts.append(f"{'#' * level} {para.text}")
        else:
            if para.text.strip():
                output_parts.append(para.text)

    for i, table in enumerate(document.tables):
        output_parts.append(f"\n[Table {i + 1}]")
        table_rows: list[list[str]] = []
        for row in table.rows:
            table_rows.append([cell.text for cell in row.cells])
        output_parts.append(json.dumps(table_rows, ensure_ascii=False))

    content = "\n".join(output_parts)
    if len(content) > _MAX_OUTPUT:
        content = content[:_MAX_OUTPUT] + "\n... (truncated)"

    return {
        "content": content,
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
    }


# --- edit ---


def _edit_com(manager: Any, resolved: str, display_path: str, **kwargs: Any) -> dict[str, Any]:
    if not os.path.isfile(resolved):
        return {"error": f"File not found: {display_path}"}

    word = manager.get_app("Word.Application")
    try:
        doc = word.Documents.Open(os.path.abspath(resolved))
    except Exception as exc:
        return {"error": f"Unable to read DOCX file: {display_path} ({type(exc).__name__}: {exc})"}

    replacements: list[dict[str, str]] = kwargs.get("replacements") or []
    content_blocks: list[dict[str, Any]] = kwargs.get("content_blocks") or []

    try:
        if not replacements and not content_blocks:
            return {"error": "Provide 'replacements' and/or 'content_blocks' for edit action"}

        if content_blocks and len(content_blocks) > _MAX_CONTENT_BLOCKS:
            return {"error": f"Too many content blocks (max {_MAX_CONTENT_BLOCKS})"}
        # NOTE: COM counts each Find.Execute iteration as one replacement.
        # The lib backend counts per-run occurrences, which may differ for
        # text spanning multiple runs. Replacement counts are approximate.
        replacements_made = 0
        for rep in replacements:
            old = rep.get("old", "")
            new = rep.get("new", "")
            if not old:
                continue
            find = doc.Content.Find
            find.ClearFormatting()
            find.Replacement.ClearFormatting()
            while find.Execute(FindText=old, ReplaceWith=new, Replace=1):
                replacements_made += 1

        rng = doc.Content
        rng.Start = rng.End
        for block in content_blocks:
            block_type = block.get("type", "paragraph")
            text = block.get("text", "")

            if block_type == "heading":
                level = block.get("level", 1)
                level = max(0, min(9, level))
                rng.InsertAfter(text + "\n")
                rng.Start = rng.End - len(text) - 1
                if level > 0:
                    rng.Style = f"Heading {level}"
                else:
                    rng.Style = "Title"
                rng.Start = rng.End
            elif block_type == "table":
                rows = block.get("rows", [])
                if not rows:
                    continue
                cols = max(len(r) for r in rows) if rows else 1
                rng.Start = rng.End
                table = doc.Tables.Add(rng, len(rows), cols)
                for i, row_data in enumerate(rows):
                    for j, cell_text in enumerate(row_data):
                        if j < cols:
                            table.Cell(i + 1, j + 1).Range.Text = str(cell_text)
                rng.Start = doc.Content.End - 1
                rng.End = rng.Start
            else:
                rng.InsertAfter(text + "\n")
                rng.Start = rng.End

        doc.Save()
    finally:
        try:
            doc.Close(SaveChanges=False)
        except Exception:
            pass

    return {
        "result": f"Edited {display_path}",
        "path": display_path,
        "replacements_made": replacements_made,
        "blocks_appended": len(content_blocks),
    }


def _edit_lib(resolved: str, display_path: str, **kwargs: Any) -> dict[str, Any]:
    import docx as _docx

    if not os.path.isfile(resolved):
        return {"error": f"File not found: {display_path}"}

    try:
        document = _docx.Document(resolved)
    except Exception as exc:
        return {"error": f"Unable to read DOCX file: {display_path} ({type(exc).__name__}: {exc})"}

    replacements: list[dict[str, str]] = kwargs.get("replacements") or []
    content_blocks: list[dict[str, Any]] = kwargs.get("content_blocks") or []

    if not replacements and not content_blocks:
        return {"error": "Provide 'replacements' and/or 'content_blocks' for edit action"}

    if content_blocks and len(content_blocks) > _MAX_CONTENT_BLOCKS:
        return {"error": f"Too many content blocks (max {_MAX_CONTENT_BLOCKS})"}

    # NOTE: lib backend counts per-run occurrences, which may differ from
    # COM backend counts (per Find.Execute iteration). Replacement counts
    # are approximate and may not match between backends.
    replacements_made = 0
    for rep in replacements:
        old = rep.get("old", "")
        new = rep.get("new", "")
        if not old:
            continue
        for para in document.paragraphs:
            if old in para.text:
                for run in para.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)
                        replacements_made += 1

    for block in content_blocks:
        block_type = block.get("type", "paragraph")
        if block_type == "heading":
            level = block.get("level", 1)
            level = max(0, min(9, level))
            document.add_heading(block.get("text", ""), level=level)
        elif block_type == "paragraph":
            document.add_paragraph(block.get("text", ""))
        elif block_type == "table":
            rows = block.get("rows", [])
            if not rows:
                continue
            cols = max(len(r) for r in rows) if rows else 1
            table = document.add_table(rows=len(rows), cols=cols)
            for i, row_data in enumerate(rows):
                for j, cell_text in enumerate(row_data):
                    if j < cols:
                        table.rows[i].cells[j].text = str(cell_text)

    document.save(resolved)
    return {
        "result": f"Edited {display_path}",
        "path": display_path,
        "replacements_made": replacements_made,
        "blocks_appended": len(content_blocks),
    }


# --- track_changes ---


def _track_changes_com(manager: Any, resolved: str, display_path: str, **kwargs: Any) -> dict[str, Any]:
    _, doc, err = _open_document_com(manager, resolved, display_path)
    if err:
        return {"error": err}

    try:
        op = kwargs.get("operation", "list")

        if op == "list":
            revisions: list[dict[str, Any]] = []
            for i in range(1, doc.Revisions.Count + 1):
                rev = doc.Revisions(i)
                revisions.append(
                    {
                        "index": i,
                        "type": str(rev.Type),
                        "author": rev.Author or "",
                        "date": str(rev.Date),
                        "text": rev.Range.Text[:200] if rev.Range else "",
                    }
                )
            return {"result": "Listed tracked changes", "revisions": revisions, "count": len(revisions)}

        if op == "accept":
            count = doc.Revisions.Count
            doc.Revisions.AcceptAll()
            doc.Save()
            return {"result": f"Accepted all {count} tracked changes", "path": display_path}

        if op == "reject":
            count = doc.Revisions.Count
            doc.Revisions.RejectAll()
            doc.Save()
            return {"result": f"Rejected all {count} tracked changes", "path": display_path}
    finally:
        try:
            doc.Close(SaveChanges=False)
        except Exception:
            pass

    return {"error": f"Unknown operation: {op}. Use 'list', 'accept', or 'reject'."}


def _track_changes_lib(resolved: str, display_path: str, **kwargs: Any) -> dict[str, Any]:
    return {"error": "Action 'track_changes' requires Windows with Office installed (COM backend)"}


# --- comments ---


def _comments_com(manager: Any, resolved: str, display_path: str, **kwargs: Any) -> dict[str, Any]:
    _, doc, err = _open_document_com(manager, resolved, display_path)
    if err:
        return {"error": err}

    try:
        op = kwargs.get("operation", "read")

        if op == "read":
            comments: list[dict[str, Any]] = []
            for i in range(1, doc.Comments.Count + 1):
                c = doc.Comments(i)
                comments.append(
                    {
                        "index": i,
                        "author": c.Author or "",
                        "date": str(c.Date),
                        "text": c.Range.Text if c.Range else "",
                        "scope": c.Scope.Text[:200] if c.Scope else "",
                    }
                )
            return {"result": "Read comments", "comments": comments, "count": len(comments)}

        if op == "add":
            comment_text = kwargs.get("comment_text")
            if not comment_text:
                return {"error": "comment_text is required for add operation"}
            rng = doc.Content
            rng.Start = rng.End
            doc.Comments.Add(Range=rng, Text=comment_text)
            doc.Save()
            return {"result": "Added comment", "path": display_path}

        if op == "delete":
            count = doc.Comments.Count
            if count == 0:
                return {"result": "No comments to delete", "path": display_path}
            for i in range(count, 0, -1):
                doc.Comments(i).Delete()
            doc.Save()
            return {"result": f"Deleted {count} comments", "path": display_path}
    finally:
        try:
            doc.Close(SaveChanges=False)
        except Exception:
            pass

    return {"error": f"Unknown operation: {op}. Use 'read', 'add', or 'delete'."}


def _comments_lib(resolved: str, display_path: str, **kwargs: Any) -> dict[str, Any]:
    return {"error": "Action 'comments' requires Windows with Office installed (COM backend)"}


# --- headers_footers ---


def _headers_footers_com(manager: Any, resolved: str, display_path: str, **kwargs: Any) -> dict[str, Any]:
    _, doc, err = _open_document_com(manager, resolved, display_path)
    if err:
        return {"error": err}

    try:
        op = kwargs.get("operation", "read")
        section_index = kwargs.get("section_index", 1)

        if section_index < 1 or section_index > doc.Sections.Count:
            return {"error": f"Section {section_index} not found. Document has {doc.Sections.Count} sections."}

        section = doc.Sections(section_index)
        # wdHeaderFooterPrimary = 1
        header_range = section.Headers(1).Range
        footer_range = section.Footers(1).Range

        if op == "read":
            return {
                "result": f"Read headers/footers for section {section_index}",
                "section": section_index,
                "header": header_range.Text.rstrip("\r\x07"),
                "footer": footer_range.Text.rstrip("\r\x07"),
            }

        if op == "set":
            header_text = kwargs.get("header_text")
            footer_text = kwargs.get("footer_text")
            if header_text is not None:
                header_range.Text = header_text
            if footer_text is not None:
                footer_range.Text = footer_text
            doc.Save()
            return {
                "result": f"Set headers/footers for section {section_index}",
                "path": display_path,
            }
    finally:
        try:
            doc.Close(SaveChanges=False)
        except Exception:
            pass

    return {"error": f"Unknown operation: {op}. Use 'read' or 'set'."}


def _headers_footers_lib(resolved: str, display_path: str, **kwargs: Any) -> dict[str, Any]:
    doc, err = _open_document_lib(resolved, display_path)
    if err:
        return {"error": err}

    op = kwargs.get("operation", "read")
    section_index = kwargs.get("section_index", 1)

    sections = list(doc.sections)
    if section_index < 1 or section_index > len(sections):
        return {"error": f"Section {section_index} not found. Document has {len(sections)} sections."}

    section = sections[section_index - 1]

    if op == "read":
        header_text = ""
        if section.header and section.header.paragraphs:
            header_text = "\n".join(p.text for p in section.header.paragraphs)
        footer_text = ""
        if section.footer and section.footer.paragraphs:
            footer_text = "\n".join(p.text for p in section.footer.paragraphs)

        return {
            "result": f"Read headers/footers for section {section_index}",
            "section": section_index,
            "header": header_text,
            "footer": footer_text,
        }

    if op == "set":
        new_header_text = kwargs.get("header_text")
        new_footer_text = kwargs.get("footer_text")

        if new_header_text is not None:
            section.header.is_linked_to_previous = False
            if section.header.paragraphs:
                section.header.paragraphs[0].text = new_header_text
            else:
                section.header.add_paragraph(new_header_text)

        if new_footer_text is not None:
            section.footer.is_linked_to_previous = False
            if section.footer.paragraphs:
                section.footer.paragraphs[0].text = new_footer_text
            else:
                section.footer.add_paragraph(new_footer_text)

        doc.save(resolved)
        return {
            "result": f"Set headers/footers for section {section_index}",
            "path": display_path,
        }

    return {"error": f"Unknown operation: {op}. Use 'read' or 'set'."}


# --- insert_image ---


def _insert_image_com(manager: Any, resolved: str, display_path: str, **kwargs: Any) -> dict[str, Any]:
    working_dir: str = kwargs.get("working_dir") or os.getcwd()
    _, doc, err = _open_document_com(manager, resolved, display_path)
    if err:
        return {"error": err}

    try:
        image_path = kwargs.get("image_path")
        if not image_path:
            return {"error": "image_path is required for insert_image"}

        img_resolved, img_err = validate_path(image_path, working_dir)
        if img_err:
            return {"error": img_err}
        if not os.path.isfile(img_resolved):
            return {"error": f"Image not found: {image_path}"}

        bookmark_name = kwargs.get("bookmark_name")
        if bookmark_name:
            if not doc.Bookmarks.Exists(bookmark_name):
                return {"error": f"Bookmark '{bookmark_name}' not found in {display_path}"}
            rng = doc.Bookmarks(bookmark_name).Range
        else:
            rng = doc.Content
            rng.Start = rng.End

        rng.InlineShapes.AddPicture(
            FileName=os.path.abspath(img_resolved),
            LinkToFile=False,
            SaveWithDocument=True,
        )

        doc.Save()
    finally:
        try:
            doc.Close(SaveChanges=False)
        except Exception:
            pass

    location = f"bookmark '{bookmark_name}'" if kwargs.get("bookmark_name") else "end of document"
    return {"result": f"Inserted image at {location} in {display_path}", "path": display_path}


def _insert_image_lib(resolved: str, display_path: str, **kwargs: Any) -> dict[str, Any]:
    from docx.shared import Inches

    working_dir: str = kwargs.get("working_dir") or os.getcwd()
    doc, err = _open_document_lib(resolved, display_path)
    if err:
        return {"error": err}

    image_path = kwargs.get("image_path")
    if not image_path:
        return {"error": "image_path is required for insert_image"}

    img_resolved, img_err = validate_path(image_path, working_dir)
    if img_err:
        return {"error": img_err}
    if not os.path.isfile(img_resolved):
        return {"error": f"Image not found: {image_path}"}

    width_inches = kwargs.get("width_inches")
    if width_inches:
        doc.add_picture(img_resolved, width=Inches(float(width_inches)))
    else:
        doc.add_picture(img_resolved)

    doc.save(resolved)
    return {"result": f"Inserted image in {display_path}", "path": display_path}


# --- styles ---


def _styles_com(manager: Any, resolved: str, display_path: str, **kwargs: Any) -> dict[str, Any]:
    _, doc, err = _open_document_com(manager, resolved, display_path)
    if err:
        return {"error": err}

    try:
        op = kwargs.get("operation", "list")

        if op == "list":
            styles: list[dict[str, str]] = []
            for i in range(1, doc.Styles.Count + 1):
                try:
                    s = doc.Styles(i)
                    # Type: 1=paragraph, 2=character, 3=table, 4=list
                    type_map = {1: "paragraph", 2: "character", 3: "table", 4: "list"}
                    styles.append(
                        {
                            "name": s.NameLocal,
                            "type": type_map.get(s.Type, "unknown"),
                            "builtin": str(s.BuiltIn),
                        }
                    )
                except Exception:
                    continue
            return {"result": "Listed styles", "styles": styles, "count": len(styles)}

        if op == "read":
            style_name = kwargs.get("style_name")
            if not style_name:
                return {"error": "style_name required for read operation"}
            try:
                s = doc.Styles(style_name)
                info = {
                    "name": s.NameLocal,
                    "builtin": str(s.BuiltIn),
                    "font_name": s.Font.Name if s.Font else "",
                    "font_size": s.Font.Size if s.Font else "",
                    "bold": str(s.Font.Bold) if s.Font else "",
                    "italic": str(s.Font.Italic) if s.Font else "",
                }
                return {"result": f"Read style '{style_name}'", "style": info}
            except Exception:
                return {"error": f"Style '{style_name}' not found"}

        if op == "set":
            style_name = kwargs.get("style_name")
            paragraph_index = kwargs.get("paragraph_index")
            if not style_name or paragraph_index is None:
                return {"error": "style_name and paragraph_index required for set operation"}
            para_idx = int(paragraph_index)
            if para_idx < 0 or para_idx >= doc.Paragraphs.Count:
                return {"error": f"paragraph_index {para_idx} out of range (0-{doc.Paragraphs.Count - 1})"}
            doc.Paragraphs(para_idx + 1).Style = style_name
            doc.Save()
            return {
                "result": f"Applied style '{style_name}' to paragraph {para_idx}",
                "path": display_path,
            }
    finally:
        try:
            doc.Close(SaveChanges=False)
        except Exception:
            pass

    return {"error": f"Unknown operation: {op}. Use 'list', 'read', or 'set'."}


def _styles_lib(resolved: str, display_path: str, **kwargs: Any) -> dict[str, Any]:
    doc, err = _open_document_lib(resolved, display_path)
    if err:
        return {"error": err}

    op = kwargs.get("operation", "list")

    if op == "list":
        styles: list[dict[str, str]] = []
        for s in doc.styles:
            type_map = {
                1: "paragraph",
                2: "character",
                3: "table",
                4: "list",
            }
            styles.append(
                {
                    "name": s.name,
                    "type": type_map.get(s.type, "unknown"),
                    "builtin": str(s.builtin),
                }
            )
        return {"result": "Listed styles", "styles": styles, "count": len(styles)}

    if op == "read":
        style_name = kwargs.get("style_name")
        if not style_name:
            return {"error": "style_name required for read operation"}
        style = None
        for s in doc.styles:
            if s.name == style_name:
                style = s
                break
        if style is None:
            return {"error": f"Style '{style_name}' not found"}
        info: dict[str, str] = {
            "name": style.name,
            "builtin": str(style.builtin),
        }
        if hasattr(style, "font") and style.font:
            info["font_name"] = str(style.font.name or "")
            info["font_size"] = str(style.font.size or "")
            info["bold"] = str(style.font.bold or "")
            info["italic"] = str(style.font.italic or "")
        return {"result": f"Read style '{style_name}'", "style": info}

    if op == "set":
        style_name = kwargs.get("style_name")
        paragraph_index = kwargs.get("paragraph_index")
        if not style_name or paragraph_index is None:
            return {"error": "style_name and paragraph_index required for set operation"}
        available_styles = {s.name for s in doc.styles}
        if style_name not in available_styles:
            return {"error": f"Style '{style_name}' not found in document"}
        para_idx = int(paragraph_index)
        if para_idx < 0 or para_idx >= len(doc.paragraphs):
            return {"error": f"paragraph_index {para_idx} out of range (0-{len(doc.paragraphs) - 1})"}
        doc.paragraphs[para_idx].style = style_name
        doc.save(resolved)
        return {
            "result": f"Applied style '{style_name}' to paragraph {para_idx}",
            "path": display_path,
        }

    return {"error": f"Unknown operation: {op}. Use 'list', 'read', or 'set'."}


# --- export_pdf ---


def _export_pdf_com(manager: Any, resolved: str, display_path: str, **kwargs: Any) -> dict[str, Any]:
    working_dir: str = kwargs.get("working_dir") or os.getcwd()
    _, doc, err = _open_document_com(manager, resolved, display_path)
    if err:
        return {"error": err}

    try:
        output_path = kwargs.get("output_path")
        if not output_path:
            output_path = os.path.splitext(resolved)[0] + ".pdf"
        else:
            out_resolved, out_err = validate_path(output_path, working_dir)
            if out_err:
                return {"error": out_err}
            output_path = out_resolved

        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        # wdExportFormatPDF = 17
        doc.ExportAsFixedFormat(
            OutputFileName=os.path.abspath(output_path),
            ExportFormat=17,
        )
    finally:
        doc.Close(SaveChanges=False)

    return {"result": f"Exported PDF to {output_path}", "path": display_path}


def _export_pdf_lib(resolved: str, display_path: str, **kwargs: Any) -> dict[str, Any]:
    return {"error": "Action 'export_pdf' requires Windows with Office installed (COM backend)"}


# --- page_setup ---


def _page_setup_com(manager: Any, resolved: str, display_path: str, **kwargs: Any) -> dict[str, Any]:
    _, doc, err = _open_document_com(manager, resolved, display_path)
    if err:
        return {"error": err}

    try:
        op = kwargs.get("operation", "set")

        if op == "read":
            ps = doc.PageSetup
            # wdOrientPortrait=0, wdOrientLandscape=1
            orient = "landscape" if ps.Orientation == 1 else "portrait"
            return {
                "result": "Read page setup",
                "orientation": orient,
                "top_margin": ps.TopMargin,
                "bottom_margin": ps.BottomMargin,
                "left_margin": ps.LeftMargin,
                "right_margin": ps.RightMargin,
                "page_width": ps.PageWidth,
                "page_height": ps.PageHeight,
                "paper_size": ps.PaperSize,
            }

        # set operation
        ps = doc.PageSetup

        orientation = kwargs.get("orientation")
        if orientation:
            # wdOrientPortrait=0, wdOrientLandscape=1
            ps.Orientation = 1 if orientation == "landscape" else 0

        paper_size = kwargs.get("paper_size")
        if paper_size:
            # wdPaperLetter=0, wdPaperA4=7, wdPaperLegal=4
            size_map = {"letter": 0, "a4": 7, "legal": 4}
            ps.PaperSize = size_map.get(paper_size, 0)

        margins = kwargs.get("margins") or {}
        # Word COM margins are in points (1 inch = 72 points)
        if margins.get("top") is not None:
            ps.TopMargin = float(margins["top"]) * 72
        if margins.get("bottom") is not None:
            ps.BottomMargin = float(margins["bottom"]) * 72
        if margins.get("left") is not None:
            ps.LeftMargin = float(margins["left"]) * 72
        if margins.get("right") is not None:
            ps.RightMargin = float(margins["right"]) * 72

        doc.Save()
    finally:
        try:
            doc.Close(SaveChanges=False)
        except Exception:
            pass

    return {"result": f"Updated page setup in {display_path}", "path": display_path}


def _page_setup_lib(resolved: str, display_path: str, **kwargs: Any) -> dict[str, Any]:
    from docx.shared import Inches

    doc, err = _open_document_lib(resolved, display_path)
    if err:
        return {"error": err}

    op = kwargs.get("operation", "set")

    section = doc.sections[0] if doc.sections else None
    if section is None:
        return {"error": "Document has no sections"}

    if op == "read":
        orient = "landscape" if section.page_width > section.page_height else "portrait"
        return {
            "result": "Read page setup",
            "orientation": orient,
            "top_margin": str(section.top_margin),
            "bottom_margin": str(section.bottom_margin),
            "left_margin": str(section.left_margin),
            "right_margin": str(section.right_margin),
            "page_width": str(section.page_width),
            "page_height": str(section.page_height),
        }

    # set operation
    orientation = kwargs.get("orientation")
    if orientation:
        if orientation == "landscape":
            if section.page_width < section.page_height:
                new_width = section.page_height
                new_height = section.page_width
                section.page_width = new_width
                section.page_height = new_height
        else:
            if section.page_width > section.page_height:
                new_width = section.page_height
                new_height = section.page_width
                section.page_width = new_width
                section.page_height = new_height

    paper_size = kwargs.get("paper_size")
    if paper_size:
        size_map = {
            "letter": (Inches(8.5), Inches(11)),
            "a4": (Inches(8.27), Inches(11.69)),
            "legal": (Inches(8.5), Inches(14)),
        }
        dims = size_map.get(paper_size)
        if dims:
            w, h = dims
            current_orient = kwargs.get("orientation", "")
            if current_orient == "landscape" or (not current_orient and section.page_width > section.page_height):
                section.page_width = h
                section.page_height = w
            else:
                section.page_width = w
                section.page_height = h

    # Margin values are provided in inches and converted via docx.shared.Inches
    margins = kwargs.get("margins") or {}
    if margins.get("top") is not None:
        section.top_margin = Inches(float(margins["top"]))
    if margins.get("bottom") is not None:
        section.bottom_margin = Inches(float(margins["bottom"]))
    if margins.get("left") is not None:
        section.left_margin = Inches(float(margins["left"]))
    if margins.get("right") is not None:
        section.right_margin = Inches(float(margins["right"]))

    doc.save(resolved)
    return {"result": f"Updated page setup in {display_path}", "path": display_path}


# --- sections ---


def _sections_com(manager: Any, resolved: str, display_path: str, **kwargs: Any) -> dict[str, Any]:
    _, doc, err = _open_document_com(manager, resolved, display_path)
    if err:
        return {"error": err}

    try:
        op = kwargs.get("operation", "list")

        if op == "list":
            sections: list[dict[str, Any]] = []
            for i in range(1, doc.Sections.Count + 1):
                s = doc.Sections(i)
                # Start types: 0=continuous, 1=new_column, 2=new_page, 3=even_page, 4=odd_page
                type_map = {
                    0: "continuous",
                    1: "new_column",
                    2: "new_page",
                    3: "even_page",
                    4: "odd_page",
                }
                sections.append(
                    {
                        "index": i,
                        "start_type": type_map.get(s.PageSetup.SectionStart, "unknown"),
                        "page_width": s.PageSetup.PageWidth,
                        "page_height": s.PageSetup.PageHeight,
                    }
                )
            return {"result": "Listed sections", "sections": sections, "count": len(sections)}

        if op == "add":
            start_type = kwargs.get("start_type", "new_page")
            # Map to COM constants
            com_type_map: dict[str, int] = {
                "continuous": 0,
                "new_column": 1,
                "new_page": 2,
                "even_page": 3,
                "odd_page": 4,
            }
            xl_type = com_type_map.get(start_type, 2)
            rng = doc.Content
            rng.Start = rng.End
            rng.InsertBreak(Type=xl_type)
            doc.Save()
            return {
                "result": f"Added {start_type} section break",
                "path": display_path,
                "total_sections": doc.Sections.Count if hasattr(doc, "Sections") else "unknown",
            }
    finally:
        try:
            doc.Close(SaveChanges=False)
        except Exception:
            pass

    return {"error": f"Unknown operation: {op}. Use 'list' or 'add'."}


def _sections_lib(resolved: str, display_path: str, **kwargs: Any) -> dict[str, Any]:
    doc, err = _open_document_lib(resolved, display_path)
    if err:
        return {"error": err}

    op = kwargs.get("operation", "list")

    if op == "list":
        sections: list[dict[str, Any]] = []
        for i, s in enumerate(doc.sections):
            start_type_raw = s.start_type if hasattr(s, "start_type") else None
            start_type_val = int(start_type_raw) if start_type_raw is not None else None
            type_map = {
                0: "continuous",
                1: "new_column",
                2: "new_page",
                3: "even_page",
                4: "odd_page",
            }
            sections.append(
                {
                    "index": i + 1,
                    "start_type": type_map.get(start_type_val, "new_page")
                    if start_type_val is not None
                    else "new_page",
                    "page_width": str(s.page_width),
                    "page_height": str(s.page_height),
                }
            )
        return {"result": "Listed sections", "sections": sections, "count": len(sections)}

    if op == "add":
        from docx.enum.section import WD_SECTION_START

        start_type = kwargs.get("start_type", "new_page")
        section_type_map: dict[str, Any] = {
            "continuous": WD_SECTION_START.CONTINUOUS,
            "new_page": WD_SECTION_START.NEW_PAGE,
            "even_page": WD_SECTION_START.EVEN_PAGE,
            "odd_page": WD_SECTION_START.ODD_PAGE,
        }
        section_start = section_type_map.get(start_type, WD_SECTION_START.NEW_PAGE)
        doc.add_section(section_start)
        doc.save(resolved)
        return {
            "result": f"Added {start_type} section break",
            "path": display_path,
            "total_sections": len(doc.sections),
        }

    return {"error": f"Unknown operation: {op}. Use 'list' or 'add'."}


# --- bookmarks ---


def _bookmarks_com(manager: Any, resolved: str, display_path: str, **kwargs: Any) -> dict[str, Any]:
    _, doc, err = _open_document_com(manager, resolved, display_path)
    if err:
        return {"error": err}

    try:
        op = kwargs.get("operation", "list")
        bookmark_name = kwargs.get("bookmark_name")

        if op == "list":
            bookmarks: list[dict[str, str]] = []
            for i in range(1, doc.Bookmarks.Count + 1):
                bm = doc.Bookmarks(i)
                bookmarks.append(
                    {
                        "name": bm.Name,
                        "start": str(bm.Start),
                        "end": str(bm.End),
                    }
                )
            return {"result": "Listed bookmarks", "bookmarks": bookmarks, "count": len(bookmarks)}

        if op == "add":
            if not bookmark_name:
                return {"error": "bookmark_name is required for add operation"}
            rng = doc.Content
            rng.Start = rng.End
            doc.Bookmarks.Add(Name=bookmark_name, Range=rng)
            doc.Save()
            return {"result": f"Added bookmark '{bookmark_name}'", "path": display_path}

        if op == "delete":
            if not bookmark_name:
                return {"error": "bookmark_name is required for delete operation"}
            if not doc.Bookmarks.Exists(bookmark_name):
                return {"error": f"Bookmark '{bookmark_name}' not found"}
            doc.Bookmarks(bookmark_name).Delete()
            doc.Save()
            return {"result": f"Deleted bookmark '{bookmark_name}'", "path": display_path}

        if op == "read":
            if not bookmark_name:
                return {"error": "bookmark_name is required for read operation"}
            if not doc.Bookmarks.Exists(bookmark_name):
                return {"error": f"Bookmark '{bookmark_name}' not found"}
            bm = doc.Bookmarks(bookmark_name)
            text = bm.Range.Text[:500] if bm.Range else ""
            return {
                "result": f"Read bookmark '{bookmark_name}'",
                "name": bm.Name,
                "text": text,
                "start": str(bm.Start),
                "end": str(bm.End),
            }
    finally:
        try:
            doc.Close(SaveChanges=False)
        except Exception:
            pass

    return {"error": f"Unknown operation: {op}. Use 'list', 'add', 'delete', or 'read'."}


def _bookmarks_lib(resolved: str, display_path: str, **kwargs: Any) -> dict[str, Any]:
    return {"error": "Action 'bookmarks' requires Windows with Office installed (COM backend)"}


# --- toc ---


def _toc_com(manager: Any, resolved: str, display_path: str, **kwargs: Any) -> dict[str, Any]:
    _, doc, err = _open_document_com(manager, resolved, display_path)
    if err:
        return {"error": err}

    try:
        op = kwargs.get("operation", "add")

        if op == "add":
            rng = doc.Content
            rng.Start = 0
            rng.End = 0
            rng.InsertBefore("\n")
            rng.Start = 0
            rng.End = 0
            doc.TablesOfContents.Add(
                Range=rng,
                UseHeadingStyles=True,
                UpperHeadingLevel=1,
                LowerHeadingLevel=3,
            )
            doc.Save()
            return {"result": "Inserted table of contents", "path": display_path}

        if op == "read":
            toc_count = doc.TablesOfContents.Count
            toc_info: list[dict[str, Any]] = []
            for i in range(1, toc_count + 1):
                toc = doc.TablesOfContents(i)
                text = toc.Range.Text[:1000] if toc.Range else ""
                toc_info.append(
                    {
                        "index": i,
                        "text": text,
                    }
                )
            return {"result": "Read table of contents", "tables_of_contents": toc_info, "count": toc_count}

        if op == "delete":
            toc_count = doc.TablesOfContents.Count
            if toc_count == 0:
                return {"result": "No table of contents to delete", "path": display_path}
            for i in range(toc_count, 0, -1):
                doc.TablesOfContents(i).Delete()
            doc.Save()
            return {"result": f"Deleted {toc_count} table(s) of contents", "path": display_path}
    finally:
        try:
            doc.Close(SaveChanges=False)
        except Exception:
            pass

    return {"error": f"Unknown operation: {op}. Use 'add', 'read', or 'delete'."}


def _toc_lib(resolved: str, display_path: str, **kwargs: Any) -> dict[str, Any]:
    return {"error": "Action 'toc' requires Windows with Office installed (COM backend)"}


# --- find_regex ---


def _find_regex_com(manager: Any, resolved: str, display_path: str, **kwargs: Any) -> dict[str, Any]:
    # NOTE: COM uses Word's wildcard syntax (not Python regex). Replacement
    # counts may differ from the lib backend which uses Python re module.
    pattern = kwargs.get("pattern")
    if not pattern:
        return {"error": "pattern is required for find_regex"}

    replace_with = kwargs.get("replace_with")
    read_only = replace_with is None

    _, doc, err = _open_document_com(manager, resolved, display_path, read_only=read_only)
    if err:
        return {"error": err}

    try:
        if replace_with is not None:
            find = doc.Content.Find
            find.ClearFormatting()
            find.Replacement.ClearFormatting()
            count = 0
            while find.Execute(
                FindText=pattern,
                MatchWildcards=True,
                ReplaceWith=replace_with,
                Replace=1,
            ):
                count += 1
            doc.Save()
            return {
                "result": f"Replaced {count} matches in {display_path}",
                "path": display_path,
                "replacements": count,
            }

        # Read-only find
        matches: list[dict[str, Any]] = []
        rng = doc.Content
        rng.Find.ClearFormatting()
        found_count = 0
        while rng.Find.Execute(FindText=pattern, MatchWildcards=True):
            matches.append(
                {
                    "text": rng.Text[:200],
                    "start": rng.Start,
                    "end": rng.End,
                }
            )
            found_count += 1
            if found_count >= 100:
                break
            rng.Start = rng.End

        return {
            "result": f"Found {len(matches)} matches",
            "matches": matches,
            "count": len(matches),
            "truncated": found_count >= 100,
        }
    finally:
        try:
            doc.Close(SaveChanges=False)
        except Exception:
            pass


def _find_regex_lib(resolved: str, display_path: str, **kwargs: Any) -> dict[str, Any]:
    doc, err = _open_document_lib(resolved, display_path)
    if err:
        return {"error": err}

    pattern = kwargs.get("pattern")
    if not pattern:
        return {"error": "pattern is required for find_regex"}

    max_pattern_len = 200
    if len(pattern) > max_pattern_len:
        return {"error": f"Regex pattern too long (max {max_pattern_len} characters)"}

    try:
        compiled = re.compile(pattern)
    except re.error as exc:
        return {"error": f"Invalid regex pattern: {exc}"}

    replace_with = kwargs.get("replace_with")

    if replace_with is not None:
        replacements = 0
        for para in doc.paragraphs:
            for run in para.runs:
                new_text, count = compiled.subn(replace_with, run.text)
                if count > 0:
                    run.text = new_text
                    replacements += count
        doc.save(resolved)
        return {
            "result": f"Replaced {replacements} matches in {display_path}",
            "path": display_path,
            "replacements": replacements,
        }

    # Read-only search
    matches: list[dict[str, Any]] = []
    found_count = 0
    for para_idx, para in enumerate(doc.paragraphs):
        for m in compiled.finditer(para.text):
            matches.append(
                {
                    "text": m.group()[:200],
                    "paragraph": para_idx,
                    "start": m.start(),
                    "end": m.end(),
                }
            )
            found_count += 1
            if found_count >= 100:
                break
        if found_count >= 100:
            break

    return {
        "result": f"Found {len(matches)} matches",
        "matches": matches,
        "count": len(matches),
        "truncated": found_count >= 100,
    }
